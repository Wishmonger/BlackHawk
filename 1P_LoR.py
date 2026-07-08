from docx import Document
import os

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Cm

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '1')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def create_Blackhawk_header(doc):
    section = doc.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=(Cm(16)))
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)
    paragraph_left = cell_left.paragraphs[0]
    paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_left = paragraph_left.add_run()
    run_left.add_picture("Blackhawk_logo.jpg")
    paragraph_right = cell_right.paragraphs[0]
    paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_right = paragraph_right.add_run(
        "\n\n1670 Riviera Ave Suite 101\nWalnut Creek, CA 94596\nT: 925.736.9990 | F: 925.984.2621")
    run_right.font.name = 'Tahoma'
    run_right.font.size = Pt(9)
    borders = OxmlElement('w:tblBorders')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')
    borders.append(bottom_border)
    table._tbl.tblPr.append(borders)

def create_1P_LoR():
    doc = Document()
    create_Blackhawk_header(doc)



    doc.save("1P_LoR.docx")


create_1P_LoR()